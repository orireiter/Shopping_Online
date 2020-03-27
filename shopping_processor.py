from flask import Flask, redirect, render_template, request
import db_for_home, db_search, barcode_maker
import datetime, os, json, email, smtplib, ssl, random
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

#this need to run everytime the page loads so adding item will be okay
def load_order(username):
    order_name = str(username)+str(datetime.date.today())+".json"
    for files in os.walk(r".\orders\\"): 
        file_checker = files[2]
    
    if order_name in file_checker:
        
        order_file = open(r".\orders\\"+str(username)+str(datetime.date.today())+".json", "r")
        order_obj = order_file.read()
        order_json = json.loads(order_obj)
        key_list = list(order_json.keys())
        value_list = list(order_json.values())
        # need to make keys - name 
        # and values - amount
        
        order_file.close()
        return value_list, key_list
        # add return to get presentation of bag, also with username
    

    else:

        order_file = open(r".\orders\\"+str(username)+str(datetime.date.today())+".json", "w")
        order_file.write("{}")
        order_file.close()
        empty = "empty"
        return empty, empty
        


# load_order("ori")

def add_item(item, username):
    path = ".\orders\\"+str(username)+str(datetime.date.today())+".json"
    
    order_file = open(path, "r")
    
    order_obj = order_file.read()
    order_json = json.loads(order_obj)
    try:
        order_json[str(item)] = str(1 + int(order_json[str(item)]))
    except:
        order_json[str(item)] = str(1)
    order_file = open(r".\orders\\"+str(username)+str(datetime.date.today())+".json", "w")
    order_json = json.dumps(order_json)
    order_file.write(order_json)
    order_file.close()
        

def delete_item(item, username):
    path = ".\orders\\"+str(username)+str(datetime.date.today())+".json"
    print(path)
    order_file1 = open(path, "r")
    print("!")
    order_obj = order_file1.read()
    order_json = json.loads(order_obj)
    
    if int(order_json[str(item)]) > 1:
        order_json[str(item)] = str(int(order_json[str(item)]) - 1)
    else:
        order_json.pop(str(item))
    order_file = open(r".\orders\\"+str(username)+str(datetime.date.today())+".json", "w")
    order_json = json.dumps(order_json)
    order_file.write(order_json)
    order_file.close()

def delete_items(item, username):
    path = ".\orders\\"+str(username)+str(datetime.date.today())+".json"
    print(path)
    order_file1 = open(path, "r")
    print("!")
    order_obj = order_file1.read()
    order_json = json.loads(order_obj)

    order_json.pop(str(item))
    order_file = open(r".\orders\\"+str(username)+str(datetime.date.today())+".json", "w")
    order_json = json.dumps(order_json)
    order_file.write(order_json)
    order_file.close()

def order_sender(username):
    order_name = str(username)+str(datetime.date.today())+".json"
    new_key_list = []
    for files in os.walk(r".\orders\\"): 
        file_checker = files[2]
    
    if order_name in file_checker:
        
        order_file = open(r".\orders\\"+str(username)+str(datetime.date.today())+".json", "r")
        order_obj = order_file.read()
        order_json = json.loads(order_obj)
        key_list = list(order_json.keys())
        value_list = list(order_json.values())
        
        order_file.close()

    else:
        print("couldnt find order")
        return "couldnt find order"
    
    # make barcodes for order
    collection = db_for_home.db_info()
    for key in key_list:
        query = collection.find_one({"Name": str(key)})
        
        if query != None:
            itemname = str(query['Name'])
            itemname2 = itemname.replace('"',"")
            itemname1 = itemname2.replace("/","")
            new_key_list.append(itemname1)
            barcode_maker.make(itemname=itemname1, barcode_num=query['Barcode'])
        else:
            new_key_list.append(key)

    # make a doc of the order
    document = Document()
    
    # making it a 2 column layout
    section = document.sections[0]
    sectPr = section._sectPr
    cols= sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')

    # defining style to be right to left
    # mystyle = document.styles.add_style('mystyle', WD_STYLE_TYPE.CHARACTER)
    
    


    # adding the information
    header =  str(datetime.date.today()) + " מזמין: " + str(username)
    document.add_heading(header, 0)
    for key in new_key_list:
        # adding info and applying bigger font
        # p = document.add_paragraph( key + " כמות: " + value_list[new_key_list.index(key)], style='List Bullet')
        p = document.add_paragraph('', style='List Bullet')
        run = p.add_run(key + " כמות: " + value_list[new_key_list.index(key)])
        font = run.font
        font.size = Pt(16)
        

        try:
            document.add_picture(r'C:\Users\Ori\Desktop\py_projects\flask\ColBo\static\barcoding\\' + key + '.jpeg', width=Cm(6))
        except:
            pass
    ran = str(random.randint(1, 10000))
    doc = r'.\static\order_docx\\' + ran + str(datetime.date.today()) + '.docx'
    document.save(doc)
    
    return doc, ran
